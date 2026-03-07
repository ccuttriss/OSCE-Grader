"""Synthetic OSCE Data Generator — LLM agent-based simulation engine.

Generates realistic end-to-end OSCE data (rubrics, student notes, faculty
scores) using independent LLM agents with unique personas for students and
faculty.  Can optionally ground generation in uploaded de-identified examples.

The key design principle is **independence**: each student agent has a unique
persona (cultural background, academic level, writing style) and each faculty
agent has a unique rater persona (experience level, scoring tendency, specialty
focus).  This prevents clustering bias and produces data that mirrors the
realistic spread seen in actual OSCE administrations.
"""

from __future__ import annotations

import io
import json
import logging
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field

import numpy as np
import pandas as pd
from openpyxl import Workbook

logger = logging.getLogger("osce_grader.synthetic")


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class SyntheticRubric:
    """A generated rubric with per-section criteria and score level descriptors."""
    assessment_type_id: str
    case_title: str
    case_description: str
    sections: dict[str, SectionRubric]
    learner_instructions: str = ""
    model_answer: str = ""
    score_table: list[tuple[str, str]] = field(default_factory=list)  # [(range, milestone), ...]


@dataclass
class SectionRubric:
    """Rubric for a single section with score-level descriptors."""
    section_key: str
    display_name: str
    max_score: int | float
    criteria: str
    score_levels: dict[int, str]  # score -> descriptor text
    checklist_items: list[dict[str, str]] = field(default_factory=list)  # [{"item": ..., "points": ..., "partial": ...}]


@dataclass
class StudentPersona:
    """A synthetic student with demographic and academic characteristics."""
    student_id: int
    name: str
    background: str
    academic_level: str
    writing_style: str
    clinical_strengths: list[str]
    clinical_weaknesses: list[str]


@dataclass
class FacultyPersona:
    """A synthetic faculty rater with scoring tendencies."""
    faculty_id: str
    name: str
    specialty: str
    years_experience: int
    scoring_tendency: str  # lenient / moderate / strict
    focus_areas: list[str]
    background_note: str


@dataclass
class SyntheticSession:
    """A fully generated session: rubric + student notes + faculty scores."""
    label: str
    assessment_type_id: str
    rubric: SyntheticRubric
    faculty: FacultyPersona
    students: list[StudentPersona]
    student_notes: dict[int, dict[str, str]]     # student_id -> {section: text}
    faculty_scores: dict[int, dict[str, float]]   # student_id -> {section: score}
    sections: list[str]


# ---------------------------------------------------------------------------
# Type metadata
# ---------------------------------------------------------------------------

_TYPE_META: dict[str, dict] = {
    "uk_osce": {
        "name": "Standard OSCE (Post-Encounter Notes)",
        "sections": {
            "hpi": ("History of Present Illness", 5),
            "pex": ("Physical Examination", 5),
            "sum": ("Summary Statement", 5),
            "ddx": ("Differential Diagnosis", 5),
            "support": ("Supporting Evidence", 5),
            "plan": ("Diagnostic Workup / Plan", 5),
        },
        "case_context": "post-encounter note for a standardized patient case",
    },
    "kpsom_ipass": {
        "name": "KPSOM I-PASS Handoff",
        "sections": {
            "illness_severity": ("Illness Severity", 2),
            "patient_summary": ("Patient Summary", 14),
            "action_list": ("Action List", 5),
            "situation_awareness": ("Situation Awareness & Contingency Planning", 3),
            "organization": ("Organization", 3),
        },
        "case_context": "I-PASS handoff note for a clinical patient encounter",
    },
    "kpsom_documentation": {
        "name": "KPSOM Clinical Documentation",
        "sections": {
            "hpi": ("History of Present Illness (PCIG)", 5),
            "social_hx": ("Social History (PCIG)", 5),
            "summary_statement": ("Summary Statement (PCDP)", 5),
            "assessment": ("Assessment (PCDP)", 5),
            "plan": ("Plan (PCDP)", 5),
            "written_communication": ("Written Communication (PCVW)", 4),
        },
        "case_context": "clinical documentation note following a patient encounter",
    },
    "kpsom_ethics": {
        "name": "KPSOM Ethics Open-Ended Questions",
        "sections": {
            "q1_total": ("Q1: Ethical Issue Identification", 4),
            "q2a_score": ("Q2A: Option Analysis - First Option", 2),
            "q2b_score": ("Q2B: Option Analysis - Second Option", 2),
            "q2c_score": ("Q2C: Option Analysis - Third Option", 2),
            "q3_total": ("Q3: Capacity Assessment Questions", 8),
        },
        "case_context": "ethics case analysis with open-ended question responses",
    },
}


# ---------------------------------------------------------------------------
# Persona generators
# ---------------------------------------------------------------------------

# Diverse backgrounds for student persona generation
_STUDENT_BACKGROUNDS = [
    "First-generation college student from a rural community in the Southern US",
    "International student from South Korea who completed pre-med in the US",
    "Career changer who previously worked as a nurse for 8 years",
    "Student from an urban underserved community with strong community health experience",
    "Second-generation immigrant from a Mexican-American family in Texas",
    "Student who grew up in West Africa and moved to the US for college",
    "Military veteran who served as a combat medic before medical school",
    "Student from an East Asian immigrant family with a research-heavy background",
    "First-generation medical student from an Appalachian community",
    "Student from a Native American reservation with IHS clinical exposure",
    "Student who grew up in the Middle East and completed undergrad in the UK",
    "Nontraditional student who worked in public health policy before med school",
    "Student from a South Asian immigrant family with dual cultural perspectives",
    "Student from a Caribbean island nation who completed pre-med in Canada",
    "First-generation Haitian-American student from an urban center",
    "Student with a PhD in biomedical engineering transitioning to clinical medicine",
    "Student from a Pacific Islander community with community health focus",
    "Eastern European immigrant student who completed undergrad in biochemistry",
    "Student from a multiracial family with extensive global health volunteer work",
    "Student who grew up in a refugee community and advocates for health equity",
]

_ACADEMIC_LEVELS = [
    "high-performing — thorough, organized, rarely misses key findings",
    "solid mid-range — covers most key points but occasionally misses nuances",
    "developing — captures main ideas but struggles with completeness and organization",
    "inconsistent — excels in some areas but has notable gaps in others",
    "strong clinically but weaker in written communication",
]

_WRITING_STYLES = [
    "concise and clinical — uses medical terminology precisely",
    "narrative and thorough — writes detailed descriptions, sometimes too verbose",
    "bullet-point oriented — structured but may lack connecting reasoning",
    "conversational tone — explains concepts clearly but sometimes informally",
    "terse and abbreviated — gets the key points across in minimal words",
]

_FACULTY_SPECIALTIES = [
    "Internal Medicine",
    "Family Medicine",
    "Emergency Medicine",
    "Pediatrics",
    "Surgery",
    "Psychiatry",
    "OB/GYN",
    "Neurology",
]


def _generate_student_personas(
    n_students: int,
    rng: np.random.Generator,
) -> list[StudentPersona]:
    """Generate diverse student personas."""
    personas = []
    backgrounds = list(_STUDENT_BACKGROUNDS)
    rng.shuffle(backgrounds)

    strengths_pool = [
        "history-taking", "physical exam documentation", "clinical reasoning",
        "differential diagnosis", "treatment planning", "patient communication",
        "organization", "summary writing", "ethical analysis", "evidence synthesis",
    ]
    weaknesses_pool = [
        "time management under pressure", "conciseness in writing",
        "including supporting evidence", "prioritizing differentials",
        "documenting negative findings", "structured organization",
        "specificity in plans", "social history completeness",
    ]

    for i in range(n_students):
        bg = backgrounds[i % len(backgrounds)]
        level = _ACADEMIC_LEVELS[int(rng.choice(len(_ACADEMIC_LEVELS)))]
        style = _WRITING_STYLES[int(rng.choice(len(_WRITING_STYLES)))]

        n_str = rng.integers(1, 4)
        n_wk = rng.integers(1, 3)
        str_indices = rng.choice(len(strengths_pool), size=n_str, replace=False)
        wk_indices = rng.choice(len(weaknesses_pool), size=n_wk, replace=False)

        personas.append(StudentPersona(
            student_id=i + 1,
            name=f"Student_{i + 1}",
            background=bg,
            academic_level=level,
            writing_style=style,
            clinical_strengths=[strengths_pool[j] for j in str_indices],
            clinical_weaknesses=[weaknesses_pool[j] for j in wk_indices],
        ))

    return personas


def _generate_faculty_persona(
    session_index: int,
    scoring_tendency: str,
    rng: np.random.Generator,
) -> FacultyPersona:
    """Generate a unique faculty rater persona for one session/year."""
    specialty = _FACULTY_SPECIALTIES[session_index % len(_FACULTY_SPECIALTIES)]
    years = int(rng.integers(3, 30))

    tendency_descriptions = {
        "lenient": "tends to give benefit of the doubt; focuses on what students got right",
        "moderate": "balanced scorer; applies rubric criteria consistently",
        "strict": "rigorous scorer; expects precise clinical language and completeness",
    }

    focus_areas_pool = [
        "clinical accuracy", "completeness of documentation",
        "organized presentation", "appropriate medical terminology",
        "evidence-based reasoning", "patient safety considerations",
        "communication clarity", "differential reasoning depth",
    ]
    n_focus = rng.integers(2, 4)
    focus_idx = rng.choice(len(focus_areas_pool), size=n_focus, replace=False)

    return FacultyPersona(
        faculty_id=f"faculty_{session_index + 1}",
        name=f"Dr. Faculty_{session_index + 1}",
        specialty=specialty,
        years_experience=years,
        scoring_tendency=scoring_tendency,
        focus_areas=[focus_areas_pool[j] for j in focus_idx],
        background_note=tendency_descriptions.get(scoring_tendency, ""),
    )


# ---------------------------------------------------------------------------
# Prompt builders
# ---------------------------------------------------------------------------

def _build_ipass_rubric_generation_prompt(
    example_rubric_text: str | None = None,
) -> list[dict[str, str]]:
    """Build prompt for generating an I-PASS handoff rubric with correct structure.

    I-PASS rubrics use checklist scoring (individual binary/partial-credit items)
    for most sections, NOT scale-based score levels.
    """
    example_block = ""
    if example_rubric_text:
        example_block = (
            "\n\nHere is a de-identified example rubric from a real administration. "
            "Use this as a template for tone, detail level, and scoring criteria "
            "structure. Generate NEW clinical content — do NOT copy verbatim:\n\n"
            f"--- EXAMPLE RUBRIC ---\n{example_rubric_text}\n--- END EXAMPLE ---\n"
        )

    system = (
        "You are a medical education assessment design expert specializing in "
        "I-PASS handoff assessments. Generate a detailed, realistic OSCE grading "
        "rubric that mirrors the structure used at KPSOM for Progress OSCEs."
    )

    user = (
        "Generate a complete KPSOM I-PASS Handoff rubric for a Progress OSCE.\n\n"
        "IMPORTANT STRUCTURAL REQUIREMENTS:\n\n"
        "1. **Peri-encounter Task Title**: A short title like "
        "'Peri-encounter Task: I-PASS Transitions of Care'\n\n"
        "2. **Learner Instructions**: Detailed paragraph(s) telling the student:\n"
        "   - The clinical situation they are in (what happened, what was done)\n"
        "   - The specific task (compose a written sign-out using I-PASS framework)\n"
        "   - Any notes about which I-PASS sections to include/omit\n\n"
        "3. **Model Answer**: A complete example of a top-scoring I-PASS handoff "
        "note, organized into the I-PASS sections (Illness Severity, Patient "
        "Summary, Action List, Situation Awareness & Contingency Planning, "
        "Synthesis by Receiver). This should be detailed and clinically accurate.\n\n"
        "4. **Case Description**: A 2-4 sentence clinical vignette describing "
        "the patient scenario.\n\n"
        "5. **Sections** — each section uses a SPECIFIC scoring format:\n\n"
        "   a) **Illness Severity** (2 pts total) — SCALE-BASED:\n"
        "      - 2 pts: Correctly identifies patient as watcher\n"
        "      - 1 pt: Indicates patient is acutely ill but uses wrong terminology\n"
        "      - 0 pts: Identifies patient as stable or omits severity\n\n"
        "   b) **Patient Summary** (14 pts total) — CHECKLIST of ~14 individual "
        "items, each worth 1 pt. Some items allow 0.5 pt partial credit. "
        "Example items from a real rubric:\n"
        "      - Age + gender: 1 pt\n"
        "      - Post-op status/presenting complaint: 1 pt\n"
        "      - PMH (full: 1 pt; partial: 0.5 pt)\n"
        "      - Key vital signs (e.g., Fever + tachycardia: 1 pt; only one: 0.5 pt)\n"
        "      - Key lab results: 1 pt each\n"
        "      - Diagnosis/clinical impression: 1 pt\n"
        "      - Treatments initiated: 1 pt each\n"
        "      The items must total exactly 14 points at full credit.\n\n"
        "   c) **Action List** (5 pts total) — CHECKLIST of ~4-5 specific action "
        "items, each worth 1 pt. Include a bonus note: 'May receive 1 pt for any "
        "other reasonable treatment/management option, not to exceed 5 pts total.'\n\n"
        "   d) **Situation Awareness & Contingency Planning** (3 pts total) — "
        "CHECKLIST of 2 specific contingency items (1 pt each) plus '1 additional "
        "pt for any other reasonable contingency plan.'\n\n"
        "   e) **Overall Organization** (3 pts total) — SCALE-BASED:\n"
        "      - Sign-out clearly separated into 4 sections (I, P, A, S): 1 pt\n"
        "      - Information organized into appropriate sections and prioritized: "
        "2 pts; minor issues: 1 pt; disorganized: 0 pts\n\n"
        "6. **Score Table**: Total score to milestone mapping. Use these EXACT "
        "milestone labels and provide 7-8 ranges covering 0–27:\n"
        "   Aspirational, Advanced Developing to Aspirational, Advanced Developing, "
        "   Mid-Developing to Advanced Developing, Mid-Developing, "
        "   Early Developing to Mid-Developing, Early Developing\n"
        f"{example_block}\n"
        "Respond with ONLY valid JSON using this EXACT structure:\n"
        "{\n"
        '  "case_title": "KPSOM I-PASS Handoff: <brief clinical title>",\n'
        '  "case_description": "<2-4 sentence clinical vignette>",\n'
        '  "learner_instructions": "<detailed instructions for the student>",\n'
        '  "model_answer": "<complete model I-PASS handoff note>",\n'
        '  "sections": {\n'
        '    "illness_severity": {\n'
        '      "criteria": "<what this section assesses>",\n'
        '      "score_levels": {"2": "...", "1": "...", "0": "..."}\n'
        "    },\n"
        '    "patient_summary": {\n'
        '      "criteria": "<what this section assesses>",\n'
        '      "checklist_items": [\n'
        '        {"item": "<clinical element>", "points": "1", "partial": null},\n'
        '        {"item": "<element with partial>", "points": "1", '
        '"partial": "<condition for 0.5 pt>"}\n'
        "      ]\n"
        "    },\n"
        '    "action_list": {\n'
        '      "criteria": "<what this section assesses>",\n'
        '      "checklist_items": [\n'
        '        {"item": "<action item>", "points": "1", "partial": null}\n'
        "      ],\n"
        '      "bonus_note": "May receive 1 pt for any other reasonable '
        'action, not to exceed 5 pts total for section"\n'
        "    },\n"
        '    "situation_awareness": {\n'
        '      "criteria": "<what this section assesses>",\n'
        '      "checklist_items": [\n'
        '        {"item": "<contingency scenario>", "points": "1", "partial": null}\n'
        "      ],\n"
        '      "bonus_note": "1 additional pt for any other reasonable '
        'contingency plan"\n'
        "    },\n"
        '    "organization": {\n'
        '      "criteria": "<what this section assesses>",\n'
        '      "score_levels": {"3": "...", "2": "...", "1": "...", "0": "..."}\n'
        "    }\n"
        "  },\n"
        '  "score_table": [\n'
        '    {"range": "26-27", "milestone": "Aspirational"},\n'
        '    {"range": "24-25", "milestone": "Advanced Developing to Aspirational"},\n'
        '    {"range": "18-23", "milestone": "Advanced Developing"},\n'
        '    {"range": "16-17", "milestone": "Mid-Developing to Advanced Developing"},\n'
        '    {"range": "14-15", "milestone": "Mid-Developing"},\n'
        '    {"range": "12-13", "milestone": "Early Developing to Mid-Developing"},\n'
        '    {"range": "0-11", "milestone": "Early Developing"}\n'
        "  ]\n"
        "}"
    )

    return [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]


def _build_ethics_rubric_generation_prompt(
    example_rubric_text: str | None = None,
) -> list[dict[str, str]]:
    """Build prompt for generating a KPSOM Ethics rubric faithful to real structure.

    The real KPSOM Ethics rubric (Case 8 pattern) has: learner instructions,
    three questions with time/point allocations, a Resources section with
    acceptable answers, a milestone score table (9 ranges, 0-18), and an
    appendix with sample capacity assessment questions.
    """
    example_block = ""
    if example_rubric_text:
        example_block = (
            "\n\nHere is a de-identified example rubric from a real administration. "
            "Use this as a template for tone, detail level, and scoring criteria "
            "structure. Generate NEW clinical content — do NOT copy verbatim:\n\n"
            f"--- EXAMPLE RUBRIC ---\n{example_rubric_text}\n--- END EXAMPLE ---\n"
        )

    system = (
        "You are a medical education assessment design expert specializing in "
        "ethics and informed consent for Progress OSCEs. Generate a detailed, "
        "realistic ethics OSCE grading rubric that mirrors the structure used at "
        "KPSOM for Progress OSCE peri-encounter ethics tasks."
    )

    user = (
        "Generate a complete KPSOM Ethics peri-encounter rubric for a Progress "
        "OSCE. The scenario must involve informed consent and decision-making "
        "capacity (e.g., a patient who may lack capacity, a surrogate decision "
        "situation, or a consent delegation issue).\n\n"
        "IMPORTANT STRUCTURAL REQUIREMENTS:\n\n"
        "1. **Case Title**: Format as 'Progress OSCE Administration X – Case Y "
        "(Patient Name)' where X is a letter (A-D) and Y is a number.\n\n"
        "2. **Case Description**: A 2-4 sentence clinical vignette describing "
        "the patient, their condition, and the ethical dilemma.\n\n"
        "3. **Learner Instructions**: A detailed paragraph structured EXACTLY "
        "like this pattern:\n"
        "   - 'Peri-encounter Task' header\n"
        "   - Reference to a preceding simulated encounter\n"
        "   - A new scenario the student must respond to\n"
        "   - '10 minutes to answer 3 questions' with sequential submission\n"
        "   - Question 1 of 3 (3 minutes / 4 points): 'Answer in 1 to 2 "
        "sentences' — two sub-parts asking the student to (a) identify what is "
        "problematic about a consent-related action and (b) how they would "
        "respond\n"
        "   - Question 2 of 3 (4 minutes / 6 points): Present EXACTLY 3 "
        "concrete options/courses of action relevant to the scenario. Ask the "
        "student to list at least one pro and one con for each option\n"
        "   - Question 3 of 3 (3 minutes / 8 points): Ask the student to write "
        "4 questions they would ask the patient to assess decision-making "
        "capacity, each addressing a different key element\n\n"
        "4. **Sections** — each section uses a SPECIFIC scoring format:\n\n"
        "   a) **q1_total** (4 pts total) — Two sub-components:\n"
        "      - 'Problematic aspects' (0-2 pts): 1 pt each for up to 2 "
        "clearly identified problems\n"
        "      - 'Reasonable response' (0-2 pts): 1 pt each for up to 2 "
        "reasonable elements of what the student would do\n"
        "      - Score levels: 4 = both sub-parts full marks, 3 = one sub-part "
        "full + one partial, 2 = one sub-part full or both partial, "
        "1 = minimal identification, 0 = nothing relevant\n\n"
        "   b) **q2a_score** (2 pts) — Option A pro/con analysis:\n"
        "      - 1 pt for at least one reasonable pro\n"
        "      - 1 pt for at least one reasonable con\n"
        "      - Score levels: 2 = both pro and con identified, 1 = only pro "
        "or only con, 0 = neither\n\n"
        "   c) **q2b_score** (2 pts) — Option B pro/con analysis (same format "
        "as Q2A)\n\n"
        "   d) **q2c_score** (2 pts) — Option C pro/con analysis (same format "
        "as Q2A)\n\n"
        "   e) **q3_total** (8 pts total) — Four capacity assessment questions:\n"
        "      - 2 pts each for questions clearly addressing one of the 5 key "
        "elements in patient-friendly language\n"
        "      - 1 pt for questions that address a key element but poorly/unclearly\n"
        "      - 0 pts for questions about orientation, diagnosis, or non-key elements\n"
        "      - No double-counting: if 2+ questions address the same element, "
        "only score the best one\n"
        "      - Score levels: 8 = all four excellent, 6 = three excellent + one "
        "partial, 4 = two good questions, 2 = one good question, 0 = none "
        "address key elements\n\n"
        "5. **Resources / Model Answer**: This is CRITICAL. Generate a complete "
        "reference answer key structured as follows:\n\n"
        "   RESOURCES\n\n"
        "   For Part A:\n"
        "   Problematic aspects of [the ethical action] include:\n"
        "   1. [List 6 specific problematic aspects relevant to the scenario]\n"
        "   ...\n"
        "   6. [Last problematic aspect]\n\n"
        "   Reasonable elements of what the student might do:\n"
        "   1. [List 3-4 reasonable responses]\n\n"
        "   For Part B:\n"
        "   Pros and cons of each option:\n"
        "   1. [Option A description]\n"
        "     - Pro: [2-3 bullet points]\n"
        "     - Con: [2-3 bullet points]\n"
        "   2. [Option B description]\n"
        "     - Pro: [2-3 bullet points]\n"
        "     - Con: [2-3 bullet points]\n"
        "   3. [Option C description]\n"
        "     - Pro: [2-3 bullet points]\n"
        "     - Con: [2-3 bullet points]\n\n"
        "   For Part C:\n"
        "   Key elements of decision-making capacity:\n"
        "   1. Appreciates current relevant medical situation\n"
        "   2. Understands intervention being proposed, with risks and benefits\n"
        "   3. Understands alternatives to proposed intervention, with risks "
        "and benefits\n"
        "   4. Is able to express or communicate a choice\n"
        "   5. Is able to give reasons for the choice\n\n"
        "   Notes:\n"
        "   - Informed consent is a process, not a document\n"
        "   - Informed consent requires a conversation\n\n"
        "   APPENDIX: Questions to Ask During an Evaluation of Medical "
        "Decision-Making Capacity\n"
        "   [Generate 3-4 sample questions for EACH of the following categories:]\n"
        "   - Questions to determine the patient's ability to understand "
        "treatment and care options\n"
        "   - Questions to determine the patient's ability to appreciate how "
        "that information applies to their situation\n"
        "   - Questions to determine the patient's ability to reason with that "
        "information\n"
        "   - Questions to determine the patient's ability to communicate and "
        "express a choice\n\n"
        "6. **Milestone Score Table**: The rubric describes what scores "
        "correspond to each developmental stage. Include a 'Mapping scores to "
        "milestones' section explaining what score ranges are expected at "
        "graduation, end of Phase 2, end of Phase 1, and entry level.\n\n"
        "Use these EXACT score table ranges:\n"
        f"{example_block}\n"
        "Respond with ONLY valid JSON using this EXACT structure:\n"
        "{\n"
        '  "case_title": "Progress OSCE Administration X – Case Y (Patient Name)",\n'
        '  "case_description": "2-4 sentence clinical vignette...",\n'
        '  "learner_instructions": "Full peri-encounter task instructions '
        'including all 3 questions with time/point allocations...",\n'
        '  "model_answer": "RESOURCES\\n\\nFor Part A:\\n...\\n\\nFor Part B:\\n'
        '...\\n\\nFor Part C:\\n...\\n\\nAPPENDIX\\n...",\n'
        '  "sections": {\n'
        '    "q1_total": {\n'
        '      "criteria": "Part A (4 possible points): 1 point each for up to '
        "two problematic aspects clearly captured in the response (2 possible "
        "points). 1 point each for up to two reasonable elements of what the "
        'student says they would do (2 possible points).",\n'
        '      "score_levels": {"4": "...", "3": "...", "2": "...", "1": "...", '
        '"0": "..."}\n'
        "    },\n"
        '    "q2a_score": {\n'
        '      "criteria": "Part B Option 1 (2 possible points): 1 point for '
        "each reasonable pro or con, maximum of 1 pro point and 1 con point. "
        'Not limited to listed examples.",\n'
        '      "score_levels": {"2": "...", "1": "...", "0": "..."}\n'
        "    },\n"
        '    "q2b_score": {\n'
        '      "criteria": "Part B Option 2 (2 possible points): same format '
        'as Option 1.",\n'
        '      "score_levels": {"2": "...", "1": "...", "0": "..."}\n'
        "    },\n"
        '    "q2c_score": {\n'
        '      "criteria": "Part B Option 3 (2 possible points): same format '
        'as Option 1.",\n'
        '      "score_levels": {"2": "...", "1": "...", "0": "..."}\n'
        "    },\n"
        '    "q3_total": {\n'
        '      "criteria": "Part C (8 possible points): 2 points for each '
        "question that clearly addresses one of the 5 key elements of "
        "decision-making capacity in clear, patient-centered language; 1 point "
        "for questions that address a key element but poorly or unclearly; no "
        "additional points if 2+ questions address the same element; no points "
        'for questions about orientation, diagnoses, etc.",\n'
        '      "score_levels": {"8": "...", "6": "...", "4": "...", "2": "...", '
        '"0": "..."}\n'
        "    }\n"
        "  },\n"
        '  "score_table": [\n'
        '    {"range": "0-3", "milestone": "Entry"},\n'
        '    {"range": "3.5-5.5", "milestone": "Entry to Early Developing"},\n'
        '    {"range": "6-7.5", "milestone": "Early Developing"},\n'
        '    {"range": "8-9.5", "milestone": "Early Developing to Mid-Developing"},\n'
        '    {"range": "10-11.5", "milestone": "Mid-Developing"},\n'
        '    {"range": "12-13", "milestone": "Mid-Developing to Advanced Developing"},\n'
        '    {"range": "13.5-15", "milestone": "Advanced Developing"},\n'
        '    {"range": "15.5-16", "milestone": "Advanced Developing to Aspirational"},\n'
        '    {"range": "16.5-18", "milestone": "Aspirational"}\n'
        "  ]\n"
        "}"
    )

    return [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]


def _build_documentation_rubric_generation_prompt(
    example_rubric_text: str | None = None,
) -> list[dict[str, str]]:
    """Build prompt for generating a KPSOM Clinical Documentation rubric.

    The real KPSOM Documentation rubric has: learner instructions with a
    pre-filled note template, milestone-based scoring (1-5) with case-specific
    criteria, competency domain tags (PCIG/PCDP/PCVW), domain-specific
    milestone mappings, and a complete model answer.
    """
    example_block = ""
    if example_rubric_text:
        example_block = (
            "\n\nHere is a de-identified example rubric from a real administration. "
            "Use this as a template for tone, detail level, and scoring criteria "
            "structure. Generate NEW clinical content — do NOT copy verbatim:\n\n"
            f"--- EXAMPLE RUBRIC ---\n{example_rubric_text}\n--- END EXAMPLE ---\n"
        )

    system = (
        "You are a medical education assessment design expert specializing in "
        "clinical documentation for Progress OSCEs. Generate a detailed, "
        "realistic clinical documentation OSCE grading rubric that mirrors "
        "the structure used at KPSOM for Progress OSCE peri-encounter tasks."
    )

    user = (
        "Generate a complete KPSOM Clinical Documentation peri-encounter "
        "rubric for a Progress OSCE.\n\n"
        "IMPORTANT STRUCTURAL REQUIREMENTS:\n\n"
        "1. **Case Title**: Format as 'Progress OSCE (Patient Name) Case X "
        "Version Y - Peri-encounter Task'\n\n"
        "2. **Case Description**: A 2-4 sentence clinical vignette describing "
        "the patient, their chief concern, and the clinical scenario.\n\n"
        "3. **Learner Instructions**: A pre-filled clinical note template that "
        "the student must complete. Structure it EXACTLY like this:\n"
        "   - 'Learner instructions:' header\n"
        "   - Instruction paragraph telling the student to complete the note "
        "by documenting HPI, social history, summary statement, assessment, "
        "and plan. State that the remainder has been provided, including "
        "physical examination findings.\n"
        "   - Pre-filled note template with:\n"
        "     * Chief Concern: (filled in with patient's chief complaint)\n"
        "     * History of Present Illness: _________ (blank for student)\n"
        "     * Birth History / Past Medical History / Past Surgical History / "
        "Allergies / Medications / Family History: (ALL filled in with "
        "case-specific details)\n"
        "     * Social History: _________ (blank for student)\n"
        "     * Physical Examination: (filled in with complete, detailed "
        "findings for all relevant systems — General, Eyes, HEENT, Cardiac, "
        "Pulmonary, Abdominal, Neurological, Extremities, etc.)\n"
        "     * Summary Statement: _________ (blank)\n"
        "     * Assessment: _________ (blank)\n"
        "     * Plan: _________ (blank)\n\n"
        "4. **Sections** — each uses MILESTONE-BASED scoring (1-5 scale) with "
        "CASE-SPECIFIC criteria. Each milestone level must list specific "
        "clinical elements the student should include, NOT generic descriptors. "
        "The criteria must reference actual clinical details from the case.\n\n"
        "   a) **hpi** (PCIG, 5 pts) — History of Present Illness:\n"
        "      - Use OPQRST-A framework (Onset, Provoking/Palliating, Quality, "
        "Region/Radiation, Severity, Timing, Associated symptoms)\n"
        "      - Each level specifies how many OPQRST elements are needed\n"
        "      - List specific pertinent positives and negatives relevant to "
        "the case (e.g., 'nausea', 'lightheadedness', 'no vomiting', "
        "'no visual changes')\n"
        "      - Entry (1): 1-2 OPQRST elements, no pertinent pos/neg\n"
        "      - Early Developing (2): 4+ OPQRST, 1-2 associated symptoms\n"
        "      - Mid-Developing (3): 6+ OPQRST, 1+ pertinent positive, "
        "1+ pertinent negative\n"
        "      - Advanced Developing (4): 6+ OPQRST, 1+ pertinent positive, "
        "2+ pertinent negatives, additional relevant history\n"
        "      - Aspirational (5): 8+ OPQRST, 2+ pertinent positives, "
        "3+ pertinent negatives, patient concerns\n\n"
        "   b) **social_hx** (PCIG, 5 pts) — Social History:\n"
        "      - Use an appropriate social history framework for the patient "
        "population (e.g., HEADSS for adolescents: Home, Education, "
        "Activities, Drugs, Sexuality, Suicide/Safety; or for adults: "
        "occupation, living situation, substance use, relationships, etc.)\n"
        "      - Each level specifies how many framework elements are needed\n"
        "      - List specific social history details relevant to the case\n"
        "      - Entry (1): May omit social history entirely\n"
        "      - Aspirational (5): Comprehensive coverage of all framework "
        "elements with case-specific details\n\n"
        "   c) **summary_statement** (PCDP, 5 pts) — Summary Statement:\n"
        "      - Each level specifies required elements (age, gender, "
        "presenting symptom, pertinent history using semantic qualifiers, "
        "relevant PE findings, pertinent negatives)\n"
        "      - List the specific semantic qualifiers relevant to the case\n"
        "      - Higher levels require more pertinent history items and "
        "pertinent negatives, with no irrelevant information\n\n"
        "   d) **assessment** (PCDP, 5 pts) — Assessment:\n"
        "      - Specify the correct primary working diagnosis for the case\n"
        "      - List the specific differential diagnoses (3-4) with their "
        "justifications/rationale based on case findings\n"
        "      - Entry (1): single diagnosis, may be incorrect, no rationale\n"
        "      - Mid-Developing (3): identifies correct primary diagnosis "
        "with rationale + 1 differential\n"
        "      - Aspirational (5): correct primary with thorough rationale + "
        "2+ differentials with reasoning for why less likely\n\n"
        "   e) **plan** (PCDP, 5 pts) — Plan:\n"
        "      - List specific next steps relevant to the case (diagnostic "
        "tests, treatments, counseling, follow-up)\n"
        "      - Specify what interventions are inappropriate for this case\n"
        "      - Higher levels include patient education, return precautions, "
        "and follow-up arrangements\n\n"
        "   f) **written_communication** (PCVW, 4 pts) — Written Communication:\n"
        "      - This is the ONE section with generic (non-case-specific) "
        "criteria about writing quality\n"
        "      - Score levels: 4 = well-organized, clear, cohesive story; "
        "3 = mostly well-organized; 2 = some disorganization; "
        "1 = significant disorganization\n"
        "      - Criteria cover: organization/cohesiveness, wordiness, "
        "clarity, typos/grammar, appropriate terminology\n\n"
        "5. **Model Answer**: Generate a COMPLETE top-scoring clinical note "
        "with ALL sections filled in. Format with section headers:\n"
        "   - HPI: (detailed, covering all OPQRST elements and pertinent "
        "pos/neg)\n"
        "   - Social History: (comprehensive coverage of the framework)\n"
        "   - Summary Statement: (concise synthesis with semantic qualifiers)\n"
        "   - Assessment: (prioritized differential with rationale)\n"
        "   - Plan: (comprehensive with diagnostics, therapeutics, education, "
        "follow-up)\n\n"
        "6. **Milestone Score Table**: Include BOTH domain-specific AND "
        "cumulative milestone mappings. Use these EXACT ranges:\n\n"
        "   PCIG (HPI + Social History = 10 points possible):\n"
        "   10 pts = Aspirational, 9 = Advanced Developing to Aspirational, "
        "8 = Advanced Developing, 7 = Mid-Developing to Advanced Developing, "
        "6 = Mid-Developing, 5 = Early Developing to Mid-Developing, "
        "4 = Early Developing, 3 = Entry to Early Developing, 2 = Entry, "
        "<2 = Behavior requiring corrective response\n\n"
        "   PCDP (Summary + Assessment + Plan = 15 points possible):\n"
        "   14-15 = Aspirational, 13 = Adv Dev to Aspirational, "
        "11-12 = Adv Dev, 10 = Mid-Dev to Adv Dev, 8-9 = Mid-Dev, "
        "7 = Early Dev to Mid-Dev, 5-6 = Early Dev, 4 = Entry to Early Dev, "
        "3 = Entry, <3 = Behavior requiring corrective response\n\n"
        "   PCVW (Written Communication = 4 points possible):\n"
        "   4 = Advanced Developing, 3.5 = Mid-Dev to Adv Dev, 3 = Mid-Dev, "
        "2.5 = Early Dev to Mid-Dev, 2 = Early Dev, 1.5 = Entry to Early Dev, "
        "1 = Entry, 0 = Behavior requiring corrective response\n\n"
        "   Cumulative (total across all 6 components = 29 points possible):\n"
        f"{example_block}\n"
        "Respond with ONLY valid JSON using this EXACT structure:\n"
        "{\n"
        '  "case_title": "Progress OSCE (Patient Name) Case X Version Y '
        '- Peri-encounter Task",\n'
        '  "case_description": "2-4 sentence clinical vignette...",\n'
        '  "learner_instructions": "Full pre-filled note template with '
        'completed sections and blanks for student to fill in...",\n'
        '  "model_answer": "HPI:\\n[detailed HPI]\\n\\nSocial History:\\n'
        '[comprehensive social history]\\n\\nSummary Statement:\\n[concise '
        'synthesis]\\n\\nAssessment:\\n[prioritized differential]\\n\\nPlan:\\n'
        '[comprehensive plan]",\n'
        '  "sections": {\n'
        '    "hpi": {\n'
        '      "criteria": "HPI (PCIG): Case-specific OPQRST criteria...",\n'
        '      "score_levels": {"5": "case-specific aspirational...", '
        '"4": "...", "3": "...", "2": "...", "1": "..."}\n'
        "    },\n"
        '    "social_hx": {\n'
        '      "criteria": "Social History (PCIG): Framework-specific...",\n'
        '      "score_levels": {"5": "...", "4": "...", "3": "...", '
        '"2": "...", "1": "..."}\n'
        "    },\n"
        '    "summary_statement": {\n'
        '      "criteria": "Summary Statement (PCDP): ...",\n'
        '      "score_levels": {"5": "...", "4": "...", "3": "...", '
        '"2": "...", "1": "..."}\n'
        "    },\n"
        '    "assessment": {\n'
        '      "criteria": "Assessment (PCDP): ...",\n'
        '      "score_levels": {"5": "...", "4": "...", "3": "...", '
        '"2": "...", "1": "..."}\n'
        "    },\n"
        '    "plan": {\n'
        '      "criteria": "Plan (PCDP): ...",\n'
        '      "score_levels": {"5": "...", "4": "...", "3": "...", '
        '"2": "...", "1": "..."}\n'
        "    },\n"
        '    "written_communication": {\n'
        '      "criteria": "Written Communication (PCVW): ...",\n'
        '      "score_levels": {"4": "...", "3": "...", "2": "...", '
        '"1": "..."}\n'
        "    }\n"
        "  },\n"
        '  "score_table": [\n'
        '    {"range": "28-29", "milestone": "Aspirational"},\n'
        '    {"range": "25-27", "milestone": "Advanced Developing to Aspirational"},\n'
        '    {"range": "23-24", "milestone": "Advanced Developing"},\n'
        '    {"range": "19-22", "milestone": "Mid-Developing to Advanced Developing"},\n'
        '    {"range": "17-18", "milestone": "Mid-Developing"},\n'
        '    {"range": "13-16", "milestone": "Early Developing to Mid-Developing"},\n'
        '    {"range": "11-12", "milestone": "Early Developing"},\n'
        '    {"range": "7-10", "milestone": "Entry to Early Developing"},\n'
        '    {"range": "6", "milestone": "Entry"}\n'
        "  ]\n"
        "}"
    )

    return [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]


def _build_rubric_generation_prompt(
    type_id: str,
    example_rubric_text: str | None = None,
) -> list[dict[str, str]]:
    """Build prompt for generating a detailed rubric."""
    # Dispatch to type-specific prompt builders
    if type_id == "kpsom_ipass":
        return _build_ipass_rubric_generation_prompt(example_rubric_text)
    if type_id == "kpsom_ethics":
        return _build_ethics_rubric_generation_prompt(example_rubric_text)
    if type_id == "kpsom_documentation":
        return _build_documentation_rubric_generation_prompt(example_rubric_text)

    meta = _TYPE_META[type_id]
    sections_desc = "\n".join(
        f"- **{display}** (max {max_s} points)"
        for _key, (display, max_s) in meta["sections"].items()
    )

    example_block = ""
    if example_rubric_text:
        example_block = (
            "\n\nHere is a de-identified example rubric from a real administration. "
            "Use this as a template for tone, detail level, and scoring criteria structure. "
            "Generate NEW content — do NOT copy verbatim:\n\n"
            f"--- EXAMPLE RUBRIC ---\n{example_rubric_text}\n--- END EXAMPLE ---\n"
        )

    system = (
        "You are a medical education assessment design expert. Generate a detailed, "
        "realistic OSCE grading rubric for the specified assessment type. The rubric "
        "must define clear, specific criteria for each score level so that different "
        "faculty members would arrive at similar scores for the same student response."
    )

    user = (
        f"Assessment type: {meta['name']}\n"
        f"Context: {meta['case_context']}\n\n"
        f"Sections to include:\n{sections_desc}\n"
        f"{example_block}\n"
        "Generate a complete rubric. For each section, provide:\n"
        "1. A brief description of what the section assesses\n"
        "2. Specific criteria at each score level (what distinguishes a 5 from a 4, "
        "a 3 from a 2, etc.)\n"
        "3. Key clinical concepts that must be present for each score level\n\n"
        "Also generate a brief clinical case vignette (2-3 sentences) that the "
        "student encountered.\n\n"
        "Respond with ONLY valid JSON:\n"
        "{\n"
        '  "case_title": "...",\n'
        '  "case_description": "2-3 sentence case vignette...",\n'
        '  "sections": {\n'
        '    "section_key": {\n'
        '      "criteria": "overall section description and grading criteria",\n'
        '      "score_levels": {"5": "descriptor...", "4": "...", ...}\n'
        "    }\n"
        "  }\n"
        "}"
    )

    return [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]


def _build_student_note_prompt(
    type_id: str,
    rubric: SyntheticRubric,
    student: StudentPersona,
    example_notes: dict[str, str] | None = None,
) -> list[dict[str, str]]:
    """Build prompt for a student agent to write their note."""
    meta = _TYPE_META[type_id]

    rubric_text = f"Case: {rubric.case_title}\n{rubric.case_description}\n\n"

    # Include learner instructions if available (tells student what to do)
    if rubric.learner_instructions:
        rubric_text += f"TASK INSTRUCTIONS:\n{rubric.learner_instructions}\n\n"

    for sec_key, sec_rubric in rubric.sections.items():
        rubric_text += (
            f"## {sec_rubric.display_name} (max {sec_rubric.max_score} pts)\n"
            f"{sec_rubric.criteria}\n\n"
        )

    example_block = ""
    if example_notes:
        example_block = (
            "\n\nHere are de-identified example student responses from a real "
            "administration. Use these as a guide for the expected FORMAT, LENGTH, "
            "and LEVEL OF DETAIL — but write entirely NEW clinical content:\n\n"
        )
        for sec, text in example_notes.items():
            example_block += f"--- {sec} ---\n{text}\n\n"

    system = (
        "You are simulating a medical student completing an OSCE post-encounter "
        "writing task. You must write AS this specific student, reflecting their "
        "unique characteristics in how they approach the task.\n\n"
        f"YOUR PERSONA:\n"
        f"- Background: {student.background}\n"
        f"- Academic level: {student.academic_level}\n"
        f"- Writing style: {student.writing_style}\n"
        f"- Clinical strengths: {', '.join(student.clinical_strengths)}\n"
        f"- Clinical weaknesses: {', '.join(student.clinical_weaknesses)}\n\n"
        "IMPORTANT: Your response quality should reflect your academic level. "
        "If you are 'developing', your notes should realistically have gaps. "
        "If you are 'high-performing', they should be thorough but still natural — "
        "not robotically perfect. Write as a real student would."
    )

    user = (
        f"You are completing a {meta['name']} assessment.\n\n"
        f"RUBRIC AND CASE:\n{rubric_text}\n"
        f"{example_block}"
        "Write your response for EACH section. Respond with ONLY valid JSON:\n"
        "{\n"
    )
    for sec_key in meta["sections"]:
        user += f'  "{sec_key}": "your response text...",\n'
    user += "}"

    return [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]


def _build_faculty_scoring_prompt(
    type_id: str,
    rubric: SyntheticRubric,
    faculty: FacultyPersona,
    student_notes: dict[str, str],
    student_persona_summary: str,
    example_scores: dict[str, float] | None = None,
) -> list[dict[str, str]]:
    """Build prompt for a faculty agent to score a student's work."""
    meta = _TYPE_META[type_id]

    rubric_text = ""
    for sec_key, sec_rubric in rubric.sections.items():
        rubric_text += (
            f"## {sec_rubric.display_name} (max {sec_rubric.max_score} pts)\n"
            f"{sec_rubric.criteria}\n"
            f"Score levels:\n"
        )
        for level, desc in sorted(sec_rubric.score_levels.items(), key=lambda x: int(x[0]), reverse=True):
            rubric_text += f"  {level}: {desc}\n"
        rubric_text += "\n"

    notes_text = ""
    for sec_key, text in student_notes.items():
        display = meta["sections"].get(sec_key, (sec_key, 0))[0]
        notes_text += f"## {display}\n{text}\n\n"

    example_block = ""
    if example_scores:
        example_block = (
            "\nFor calibration reference, here are scores from a prior administration "
            "for a student of similar performance level (use as a general anchor, not "
            "an exact guide):\n"
        )
        for sec, score in example_scores.items():
            example_block += f"  {sec}: {score}\n"

    system = (
        "You are a medical faculty member grading a student's OSCE response. "
        "Score each section independently using the rubric criteria.\n\n"
        f"YOUR RATER PROFILE:\n"
        f"- Name: {faculty.name}\n"
        f"- Specialty: {faculty.specialty}\n"
        f"- Experience: {faculty.years_experience} years of clinical teaching\n"
        f"- Scoring tendency: {faculty.scoring_tendency} — {faculty.background_note}\n"
        f"- Focus areas: {', '.join(faculty.focus_areas)}\n\n"
        "IMPORTANT: Score realistically as this faculty member would. Your scoring "
        "tendency should subtly influence your scores — a lenient rater gives more "
        "benefit of the doubt, a strict rater expects precise completeness. "
        "But still follow the rubric — your tendency is a ~0.5 point bias, not a "
        "wholesale departure from the criteria."
    )

    user = (
        f"RUBRIC:\n{rubric_text}\n"
        f"STUDENT RESPONSE:\n{notes_text}\n"
        f"{example_block}\n"
        "Score each section. Respond with ONLY valid JSON — section keys mapped "
        "to numeric scores:\n"
        "{\n"
    )
    for sec_key, (_, max_s) in meta["sections"].items():
        user += f'  "{sec_key}": <score 0-{max_s}>,\n'
    user += "}"

    return [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]


# ---------------------------------------------------------------------------
# JSON parsing helpers
# ---------------------------------------------------------------------------

def _strip_fences(text: str) -> str:
    """Strip markdown code fences from LLM output."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n", 1)
        text = lines[1] if len(lines) > 1 else ""
    if text.endswith("```"):
        text = text[:-3]
    return text.strip()


def _parse_json_response(text: str) -> dict:
    """Parse a JSON response from an LLM, handling common quirks."""
    text = _strip_fences(text)
    # Try direct parse
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    # Try extracting JSON object from surrounding text
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            pass
    raise ValueError(f"Could not parse JSON from LLM response:\n{text[:500]}")


def _parse_rubric_response(text: str, type_id: str) -> SyntheticRubric:
    """Parse LLM rubric generation response into a SyntheticRubric."""
    data = _parse_json_response(text)
    meta = _TYPE_META[type_id]

    sections = {}
    raw_sections = data.get("sections", {})
    for sec_key, (display, max_s) in meta["sections"].items():
        sec_data = raw_sections.get(sec_key, {})

        # Parse score_levels (scale-based sections)
        score_levels = {}
        raw_levels = sec_data.get("score_levels", {})
        for k, v in raw_levels.items():
            try:
                score_levels[int(k)] = str(v)
            except (ValueError, TypeError):
                pass

        # Parse checklist_items (checklist-based sections like Patient Summary)
        checklist_items = []
        raw_items = sec_data.get("checklist_items", [])
        for item in raw_items:
            if isinstance(item, dict):
                checklist_items.append({
                    "item": str(item.get("item", "")),
                    "points": str(item.get("points", "1")),
                    "partial": str(item["partial"]) if item.get("partial") else None,
                })

        # Append bonus_note to criteria if present
        criteria = sec_data.get("criteria", "")
        bonus_note = sec_data.get("bonus_note")
        if bonus_note:
            criteria = f"{criteria}\n({bonus_note})"

        sections[sec_key] = SectionRubric(
            section_key=sec_key,
            display_name=display,
            max_score=max_s,
            criteria=criteria,
            score_levels=score_levels,
            checklist_items=checklist_items,
        )

    # Parse score_table
    score_table = []
    raw_table = data.get("score_table", [])
    for entry in raw_table:
        if isinstance(entry, dict):
            score_table.append((
                str(entry.get("range", "")),
                str(entry.get("milestone", "")),
            ))

    return SyntheticRubric(
        assessment_type_id=type_id,
        case_title=data.get("case_title", "Synthetic Case"),
        case_description=data.get("case_description", ""),
        sections=sections,
        learner_instructions=data.get("learner_instructions", ""),
        model_answer=data.get("model_answer", ""),
        score_table=score_table,
    )


# ---------------------------------------------------------------------------
# Session generation orchestrator
# ---------------------------------------------------------------------------

def generate_synthetic_session(
    type_id: str,
    session_index: int,
    n_students: int,
    variability: float,
    llm_caller,
    *,
    temperature: float = 0.7,
    example_rubric_text: str | None = None,
    example_notes: dict[str, str] | None = None,
    example_scores: dict[int, dict[str, float]] | None = None,
    seed: int = 42,
    progress_callback=None,
) -> SyntheticSession:
    """Generate one complete synthetic session using LLM agents.

    Parameters
    ----------
    type_id : str
        Assessment type ID.
    session_index : int
        Index of this session (0-based), used for faculty persona variation.
    n_students : int
        Number of student agents to generate.
    variability : float
        0.0–1.0 controlling scoring tendency distribution and temperature.
    llm_caller : callable
        ``(messages, temperature, top_p) -> str`` LLM call function.
    temperature : float
        Base LLM temperature for generation.
    example_rubric_text : str, optional
        De-identified rubric text to ground generation.
    example_notes : dict, optional
        Example student notes ``{section: text}`` for grounding.
    example_scores : dict, optional
        Example faculty scores ``{student_id: {section: score}}`` for calibration.
    seed : int
        Random seed for persona generation.
    progress_callback : callable, optional
        ``(step_name: str, current: int, total: int) -> None``
    """
    meta = _TYPE_META.get(type_id)
    if meta is None:
        raise ValueError(f"Unknown assessment type: {type_id}")

    rng = np.random.default_rng(seed + session_index * 1000)
    sections = list(meta["sections"].keys())

    # Effective temperature scales with variability
    eff_temp = temperature * (0.5 + variability)

    # --- Step 1: Generate rubric ---
    if progress_callback:
        progress_callback("Generating rubric", 0, n_students + 2)

    rubric_messages = _build_rubric_generation_prompt(type_id, example_rubric_text)
    rubric_response = llm_caller(rubric_messages, eff_temp, 1.0)
    rubric = _parse_rubric_response(rubric_response, type_id)

    # --- Step 2: Generate faculty persona ---
    tendencies = ["lenient", "moderate", "strict"]
    # Variability affects the distribution: low var = mostly moderate
    if variability < 0.3:
        tendency = "moderate"
    elif variability < 0.7:
        tendency = tendencies[session_index % len(tendencies)]
    else:
        # High variability: more extreme tendencies
        weights = [0.35, 0.30, 0.35]
        tendency = rng.choice(tendencies, p=weights)

    faculty = _generate_faculty_persona(session_index, tendency, rng)

    # --- Step 3: Generate student personas ---
    students = _generate_student_personas(n_students, rng)

    # --- Step 4: Generate student notes (parallel) ---
    student_notes: dict[int, dict[str, str]] = {}
    # Pick one example student's scores for calibration reference if available
    example_score_ref = None
    if example_scores:
        first_key = next(iter(example_scores))
        example_score_ref = example_scores[first_key]

    if progress_callback:
        progress_callback("Generating student notes (parallel)", 1, n_students + 2)

    def _generate_one_note(student: StudentPersona) -> tuple[int, dict[str, str]]:
        note_messages = _build_student_note_prompt(
            type_id, rubric, student, example_notes
        )
        note_response = llm_caller(note_messages, eff_temp, 1.0)
        try:
            notes = _parse_json_response(note_response)
            for sec in sections:
                if sec not in notes:
                    notes[sec] = ""
            return student.student_id, {
                sec: str(notes.get(sec, "")) for sec in sections
            }
        except ValueError:
            logger.warning(
                "Failed to parse student %d notes, using empty.",
                student.student_id,
            )
            return student.student_id, {sec: "" for sec in sections}

    with ThreadPoolExecutor(max_workers=min(n_students, 8)) as pool:
        note_futures = {
            pool.submit(_generate_one_note, student): student
            for student in students
        }
        notes_done = 0
        for future in as_completed(note_futures):
            sid, parsed_notes = future.result()
            student_notes[sid] = parsed_notes
            notes_done += 1
            if progress_callback:
                progress_callback(
                    f"Student notes ({notes_done}/{n_students})",
                    notes_done,
                    n_students + 2,
                )

    # --- Step 5: Faculty scores all students (parallel) ---
    faculty_scores: dict[int, dict[str, float]] = {}

    if progress_callback:
        progress_callback("Scoring students (parallel)", n_students, n_students + 2)

    def _score_one_student(student: StudentPersona) -> tuple[int, dict[str, float]]:
        score_messages = _build_faculty_scoring_prompt(
            type_id,
            rubric,
            faculty,
            student_notes[student.student_id],
            f"Level: {student.academic_level}",
            example_score_ref,
        )
        score_response = llm_caller(score_messages, max(0.2, eff_temp * 0.5), 1.0)
        try:
            raw_scores = _parse_json_response(score_response)
            parsed: dict[str, float] = {}
            for sec in sections:
                val = raw_scores.get(sec)
                if val is not None:
                    try:
                        parsed[sec] = float(val)
                    except (ValueError, TypeError):
                        parsed[sec] = 0.0
                else:
                    parsed[sec] = 0.0
            return student.student_id, parsed
        except ValueError:
            logger.warning(
                "Failed to parse faculty scores for student %d, using zeros.",
                student.student_id,
            )
            return student.student_id, {sec: 0.0 for sec in sections}

    with ThreadPoolExecutor(max_workers=min(n_students, 8)) as pool:
        score_futures = {
            pool.submit(_score_one_student, student): student
            for student in students
        }
        scores_done = 0
        for future in as_completed(score_futures):
            sid, parsed_scores = future.result()
            faculty_scores[sid] = parsed_scores
            scores_done += 1
            if progress_callback:
                progress_callback(
                    f"Faculty scoring ({scores_done}/{n_students})",
                    n_students + 1,
                    n_students + 2,
                )

    label = f"{2020 + session_index} {'Spring' if session_index % 2 == 0 else 'Fall'}"

    return SyntheticSession(
        label=label,
        assessment_type_id=type_id,
        rubric=rubric,
        faculty=faculty,
        students=students,
        student_notes=student_notes,
        faculty_scores=faculty_scores,
        sections=sections,
    )


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------

def rubric_to_display_text(rubric: SyntheticRubric) -> str:
    """Format a rubric as readable text for UI display."""
    lines = [
        f"# {rubric.case_title}",
        f"\n{rubric.case_description}\n",
    ]

    # Learner instructions
    if rubric.learner_instructions:
        lines.append("\n## Learner Instructions")
        lines.append(rubric.learner_instructions)
        lines.append("")

    # Model answer
    if rubric.model_answer:
        lines.append("\n## Model Answer")
        lines.append(rubric.model_answer)
        lines.append("")

    # Grading rubric sections
    lines.append("\n## Grading Rubric\n")
    for sec_key, sec in rubric.sections.items():
        lines.append(f"\n### {sec.display_name} ({sec.max_score} pts total)")
        lines.append(sec.criteria)

        # Checklist-based sections (e.g., Patient Summary, Action List)
        if sec.checklist_items:
            lines.append("")
            for item in sec.checklist_items:
                partial = item.get("partial")
                if partial:
                    lines.append(
                        f"- {item['item']}: {item['points']} pt; {partial}"
                    )
                else:
                    lines.append(f"- {item['item']}: {item['points']} pt")

        # Scale-based sections (e.g., Illness Severity, Organization)
        if sec.score_levels:
            lines.append("\n**Score Levels:**\n")
            for level in sorted(sec.score_levels.keys(), reverse=True):
                lines.append(f"**{level}:** {sec.score_levels[level]}\n")

    # Score table
    if rubric.score_table:
        lines.append("\n## Score Table\n")
        lines.append("| Total Score | Milestone |")
        lines.append("|-------------|-----------|")
        for score_range, milestone in rubric.score_table:
            lines.append(f"| {score_range} | {milestone} |")

    return "\n".join(lines)


def rubric_to_excel(rubric: SyntheticRubric, type_id: str) -> bytes:
    """Export rubric as an Excel file matching the expected input format."""
    meta = _TYPE_META[type_id]
    wb = Workbook()
    ws = wb.active
    ws.title = "Rubric"

    if type_id == "uk_osce":
        # UK OSCE expects rubric.xlsx with section columns and one data row
        sections = list(meta["sections"].keys())
        ws.append(sections)
        row = []
        for sec in sections:
            sec_rubric = rubric.sections.get(sec)
            if sec_rubric:
                text = sec_rubric.criteria
                if sec_rubric.score_levels:
                    text += "\n" + "\n".join(
                        f"{lvl}) {desc}"
                        for lvl, desc in sorted(sec_rubric.score_levels.items(), reverse=True)
                    )
                row.append(text)
            else:
                row.append("")
        ws.append(row)
    else:
        # KPSOM types: .docx-style rubric exported as structured Excel
        # Add learner instructions if present
        if rubric.learner_instructions:
            ws.append(["Learner Instructions"])
            ws.append([rubric.learner_instructions])
            ws.append([])  # blank row separator

        # Add model answer if present
        if rubric.model_answer:
            ws.append(["Model Answer"])
            ws.append([rubric.model_answer])
            ws.append([])  # blank row separator

        ws.append(["Section", "Criteria", "Max Score", "Score Levels / Checklist"])
        for sec_key, sec in rubric.sections.items():
            # Build scoring details text
            if sec.checklist_items:
                # Checklist-based section
                items_lines = []
                for item in sec.checklist_items:
                    partial = item.get("partial")
                    if partial:
                        items_lines.append(
                            f"{item['item']}: {item['points']} pt; {partial}"
                        )
                    else:
                        items_lines.append(f"{item['item']}: {item['points']} pt")
                scoring_text = "\n".join(items_lines)
            elif sec.score_levels:
                # Scale-based section
                scoring_text = "\n".join(
                    f"{lvl}: {desc}"
                    for lvl, desc in sorted(sec.score_levels.items(), reverse=True)
                )
            else:
                scoring_text = ""
            ws.append([sec.display_name, sec.criteria, sec.max_score, scoring_text])

        # Add score table if present
        if rubric.score_table:
            ws.append([])  # blank row separator
            ws.append(["Total Score", "Milestone"])
            for score_range, milestone in rubric.score_table:
                ws.append([score_range, milestone])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def rubric_to_docx(rubric: SyntheticRubric) -> bytes:
    """Export rubric as a .docx file suitable for the KPSOM parsing pipeline."""
    from docx import Document

    doc = Document()
    doc.add_heading(rubric.case_title or "Rubric", level=1)

    if rubric.learner_instructions:
        doc.add_heading("Learner Instructions", level=2)
        doc.add_paragraph(rubric.learner_instructions)

    for sec_key, sec in rubric.sections.items():
        doc.add_heading(sec.display_name, level=2)
        doc.add_paragraph(f"Max Score: {sec.max_score}")
        if sec.criteria:
            doc.add_paragraph(sec.criteria)
        if sec.checklist_items:
            for item in sec.checklist_items:
                partial = item.get("partial", "")
                line = f"{item['item']}: {item['points']} pt"
                if partial:
                    line += f"; {partial}"
                doc.add_paragraph(line, style="List Bullet")
        elif sec.score_levels:
            for lvl, desc in sorted(sec.score_levels.items(), reverse=True):
                doc.add_paragraph(f"{lvl}: {desc}", style="List Bullet")

    if rubric.model_answer:
        doc.add_heading("Model Answer", level=2)
        doc.add_paragraph(rubric.model_answer)

    if rubric.score_table:
        doc.add_heading("Score Table", level=2)
        for score_range, milestone in rubric.score_table:
            doc.add_paragraph(f"{score_range}: {milestone}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def answer_key_to_excel(rubric: SyntheticRubric, type_id: str) -> bytes:
    """Export an answer key derived from the rubric's top-level criteria."""
    meta = _TYPE_META[type_id]
    wb = Workbook()
    ws = wb.active
    ws.title = "Answer Key"

    sections = list(meta["sections"].keys())
    ws.append(sections)
    row = []
    for sec in sections:
        sec_rubric = rubric.sections.get(sec)
        if sec_rubric and sec_rubric.score_levels:
            # Use the highest score level as the answer key
            max_level = max(sec_rubric.score_levels.keys())
            row.append(sec_rubric.score_levels[max_level])
        else:
            row.append("")
    ws.append(row)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def student_notes_to_excel(session: SyntheticSession) -> bytes:
    """Export student notes as an Excel file matching the expected input format."""
    type_id = session.assessment_type_id
    meta = _TYPE_META[type_id]
    wb = Workbook()
    ws = wb.active
    ws.title = "Student Notes"

    sections = session.sections

    if type_id in ("kpsom_ipass", "kpsom_documentation"):
        # KPSOM format: row 0 = Q-numbers, row 1 = section headers, rows 2+ = data
        q_row = [""] + [f"Q{i+1}" for i in range(len(sections))]
        ws.append(q_row)
        header_row = ["Student"] + [
            meta["sections"][s][0] for s in sections
        ]
        ws.append(header_row)
        for sid in sorted(session.student_notes.keys()):
            row = [sid] + [session.student_notes[sid].get(s, "") for s in sections]
            ws.append(row)
    elif type_id == "kpsom_ethics":
        # Ethics: row 0 = question headers, rows 1+ = data
        header_row = ["Student"] + [
            meta["sections"][s][0] for s in sections
        ]
        ws.append(header_row)
        for sid in sorted(session.student_notes.keys()):
            row = [sid] + [session.student_notes[sid].get(s, "") for s in sections]
            ws.append(row)
    else:
        # UK OSCE: simple flat format
        headers = sections
        ws.append(headers)
        for sid in sorted(session.student_notes.keys()):
            row = [session.student_notes[sid].get(s, "") for s in sections]
            ws.append(row)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def faculty_scores_to_excel(session: SyntheticSession) -> bytes:
    """Export faculty scores matching the format load_faculty_session() expects."""
    # Reuse session_to_excel from gold_standard by constructing a SessionData
    from gold_standard import SessionData, session_to_excel

    sd = SessionData(
        label=session.label,
        assessment_type_id=session.assessment_type_id,
        sections=session.sections,
        scores=session.faculty_scores,
        student_count=len(session.faculty_scores),
    )
    return session_to_excel(sd)


def session_to_zip(session: SyntheticSession) -> bytes:
    """Bundle an entire session into a downloadable ZIP."""
    import zipfile

    buf = io.BytesIO()
    safe_label = session.label.replace(" ", "_")
    type_id = session.assessment_type_id

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # Rubric
        zf.writestr(
            f"{safe_label}_rubric.xlsx",
            rubric_to_excel(session.rubric, type_id),
        )
        # Answer key (UK OSCE) or rubric detail
        if type_id == "uk_osce":
            zf.writestr(
                f"{safe_label}_answer_key.xlsx",
                answer_key_to_excel(session.rubric, type_id),
            )
        # Student notes
        zf.writestr(
            f"{safe_label}_student_notes.xlsx",
            student_notes_to_excel(session),
        )
        # Faculty scores
        zf.writestr(
            f"{safe_label}_faculty_scores.xlsx",
            faculty_scores_to_excel(session),
        )
        # Rubric as readable text
        zf.writestr(
            f"{safe_label}_rubric.txt",
            rubric_to_display_text(session.rubric),
        )

    return buf.getvalue()


def all_sessions_to_zip(sessions: list[SyntheticSession]) -> bytes:
    """Bundle multiple sessions into a single ZIP."""
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for session in sessions:
            safe_label = session.label.replace(" ", "_")
            type_id = session.assessment_type_id
            prefix = f"{safe_label}/"

            zf.writestr(
                f"{prefix}rubric.xlsx",
                rubric_to_excel(session.rubric, type_id),
            )
            if type_id == "uk_osce":
                zf.writestr(
                    f"{prefix}answer_key.xlsx",
                    answer_key_to_excel(session.rubric, type_id),
                )
            zf.writestr(
                f"{prefix}student_notes.xlsx",
                student_notes_to_excel(session),
            )
            zf.writestr(
                f"{prefix}faculty_scores.xlsx",
                faculty_scores_to_excel(session),
            )
            zf.writestr(
                f"{prefix}rubric.txt",
                rubric_to_display_text(session.rubric),
            )

    return buf.getvalue()
